############################################################################################################
# Verification of SUMMARY_MEANS and part_processed hourly files
# Author: CAS
# Date: 05/09/2024
# Version: 1.2
# 1.2 Verification check merged together with data QC checks.
# 1.1 Edited to output verification checks in word document rather than in excel work book as in version 1.0
############################################################################################################
# IMPORTING PACKAGES #
import docx
#import config
import os
import pandas as pd
from docx.shared import RGBColor
import operator
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from wavepostprocessing.Housekeeping import filenames_to_remove
from wavepostprocessing.config import load_config, print_message
#from Housekeeping import filenames_to_remove
#from config import load_config, print_message
import sys
from datetime import datetime

# --- Creating verification log --- #
def create_verif_log(log_header):
    """
    This function creates a verification log as a docx format.
    :return: verif_log: Returning a verification log, where verification checks of the SUMMARY_MEANS and part_processed hourly files will be outputted
    """
    verif_log = docx.Document()
    verif_log.add_heading(f"{log_header} - {config.get('pc_date')}")
    return verif_log

def add_header(log_header):
    """
    Adding header to verification.
    :param log_header: Header in verification log.
    :return:
    """
    if len(verif_log.paragraphs) > 2:
        verif_log.add_page_break()
    verif_log.add_heading(f"{log_header} - {config.get('pc_date')}")
    save_verif_log(verif_log)

def save_verif_log(verif_log):
    """
    Function to save the verification log in the results folder.
    :param verif_log: The verification log
    :return:
    """
    verif_log_file_path = os.path.join(config.get('root_folder'), config.get('log_folder'), f"{config.get('verif_name')}_{config.get('pc_date')}.docx")
    verif_log.save(verif_log_file_path)


# --- Adding text to log --- #
def add_text(log, text_to_log, x, y, z):
    """
    This function adds text to the verification log
    :param log: The verification log that the output gets printed to.
    :param text_to_log: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_to_log}")
    run.bold = True
    run.font.color.rgb = RGBColor(x, y, z)

# --- Adding description text to log --- #
def add_description_text(log, description, x, y, z):
    """
    This function adds a description text to the verification log
    :param log: The verification log that the output gets printed to.
    :param description: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{description}")
    #run.bold = False
    run.font.color.rgb = RGBColor(x, y, z)

def add_text_no_error(log, text_no_error):
    """
    This function adds text to the verification log if no errors and prints it in green.
    :param log: The verification log that the output gets printed to.
    :param text_no_error: The explanatory text that are being printed to the log.
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_no_error}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 155, 0)


# --- Changing docx to landscape mode --- #
def landscape(log):
    """
    Changing the docx to landscape mode
    :param log: The document that is being changed.
    :return:
    """
    new_section = log.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

# --- Changing docx to portrait mode --- #
def portrait(log):
    """ Changing the docx to portrait mode
    :param log: The document that is being changed
    :return:
    """
    new_section = log.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

# --- CHECKING IF DATASET EXISTS AND THEN READING IT IN --- #
def dataframe(file_name, variable):
    """
    Importing dataset as dataframe and creating a flag if the dataset doesn't exist.
    :param file_name: The dataset to import.
    :return: df. The dataset as dataframe if it exists.
    :return: file_exists. Flag to indicate if the dataset exists.
    """
    dataframe_path = os.path.join(config.get('root_folder'), config.get('results_folder'), config.get('summary_folder'), f'{file_name}.csv')
    if os.path.exists(dataframe_path):
        df = pd.read_csv(dataframe_path, dtype={'subject_code': str})
        file_exists = True
        if config.get('run_housekeeping').lower() == 'yes':
            df = df[(~df[variable].isin(filenames_to_remove))]

    if not os.path.exists(dataframe_path):
        df = None
        file_exists = False
    return df, file_exists


# --- PRINTING FILE AND DEVICE INFORMATION TO VERIFICATION LOG --- #
def information_to_verif_log(log, df, table, variable_to_count, text_to_log, count_mode):
    """
    This function counts the frequency of a variable in the summary dataset and prints this together with a table to the verification log.
    :param log: the verification log that it is being printed to.
    :param variable_to_count: The variable which is being counted and printed in the qc log.
    :param text_to_log: The text that will be printed above the table in the data qc log.
    :param count_mode: Defining if all should be counted or only unique values of the variable should be counted.
    :return:
    """
    # Counting frequency
    if count_mode == 'total':
        count = df[variable_to_count].count()
    elif count_mode == 'unique':
        count = df[variable_to_count].nunique()

    # Adding to the qc_log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_to_log}: {count}")
    run.bold = True

    if table == 'Yes':
        # Calculating frequency of the variable
        frequency = df[variable_to_count].value_counts()

        # Adding table to qc log
        table = log.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Defining the header row:
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Frequency'
        hdr_cells[1].text = variable_to_count

        # Adding the values to the table
        for file_value, freq in frequency.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(freq)
            row_cells[1].text = str(file_value)

        log.add_paragraph("\n")
        save_verif_log(log)

    else:
        save_verif_log(log)

# --- GENERATING INCLUDE CRITERIA --- ~
def include_criteria(log, df, text_to_log, description):
    """
    This function creates an include variable, to tag if data files should be included in release based on certain criteria.
    :param log: Where the verification checks gets outputted to.
    :param df: Summary_df = The summary overall mean as a dataframe.
    :param text_to_log: The text printed to the verification log to show what is being checked.
    :param x, y, z: Specifies what color the text printed to the log should be.
    :return: df: Returning the updated dataframe.
    """
    # Generating include criteria
    df['include'] = 0
    df.loc[
        (df['Pwear'] > config.get('ver_pwear')) &
        (df['Pwear'].notnull()) &
        (df['Pwear_morning'] >= config.get('ver_pwear_morn')) &
        (df['Pwear_noon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_afternoon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_night'] >= config.get('ver_pwear_quad')), 'include'] = 1

    if config.get('impute_data').lower() == 'yes':
        df.loc[
            (df['Pwear'] > config.get('ver_pwear')) &
            (df['Pwear'].notnull()) &
            (df['Pwear_morning'] < config.get('ver_pwear_morn')) &
            (df['Pwear_noon'] >= config.get('ver_pwear_quad')) &
            (df['Pwear_afternoon'] >= config.get('ver_pwear_quad')) &
            (df['Pwear_night'] >= config.get('ver_pwear_quad')) &
            (df['include'] != 1), 'include'] = 2


    # Summarising include variable
    include_freq = df['include'].value_counts()
    include_percent = df['include'].value_counts(normalize=True) * 100
    include_cum_percent = include_percent.cumsum()

    # Adding to the verification log
    add_text(log, text_to_log, 0, 0, 0)
    add_description_text(log, description, 0, 0, 0)
    if config.get('impute_data').lower() == 'yes':
        add_description_text(log, f"Include = 2: Sleep data are imputed (Pwear >= {config.get('ver_pwear')}, Pwear morning < {config.get('ver_pwear_morn')} and Pwear noon/afternoon/night >= {config.get('ver_pwear_quad')}).", 0, 0, 0)

    # Adding table to verification log
    table = log.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Include', 'Frequency', 'Percent', 'Cumulative percent']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    # Adding the values to the table
    for idx, value in include_freq.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(value)
        row_cells[2].text = f"{round(include_percent[idx], 0)}"
        row_cells[3].text = f"{round(include_cum_percent[idx], 0)}"

    total_frequency = round(include_freq.sum(), 0)
    total_percent = round(include_percent.sum(), 0)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Total'
    row_cells[1].text = str(total_frequency)
    row_cells[2].text = f"{total_percent}"
    row_cells[3].text = ""

    log.add_paragraph("\n")
    save_verif_log(log)

    return df

# --- Summarising startdate to check it falls inside the expected testing dates --- #
def sum_startdate(log, df, text_to_log, description, x, y, z):
    """
    Summarising the start time variable to be able to check if they all fall in between the expected test dates.
    :param log: Verification log where the verification checks gets outputted to.
    :param df: Summary_df = The summary overall mean as a dataframe.
    :param text_to_log: The text printed to the verification log to show what is being checked.
    :param x, y, z: Specifies the color that the text will be printed in
    :return: df: Returning the updated dataframe.
    """

    # Converting startdate to datetime format
    #df['startdate'] = pd.to_datetime(df['startdate'], format='%d/%m/%Y')
    df['startdate'] = pd.to_datetime(df['startdate'], format='%Y-%m-%d', errors='coerce')

    # Adding to the verification log
    add_text(log, text_to_log, x, y, z)
    add_description_text(log, description, x, y, z)

    # Adding table to verification log
    table = log.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Variable', 'Obs', 'Mean', 'Min', 'Max']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    # Adding values to the table
    row_cells = table.add_row().cells
    row_cells[0].text = "Start date"
    row_cells[1].text = str(df['startdate'].count())
    row_cells[2].text = str(df['startdate'].mean().strftime('%d%b%Y'))
    row_cells[3].text = str(df['startdate'].min().strftime('%d%b%Y'))
    row_cells[4].text = str(df['startdate'].max().strftime('%d%b%Y'))

    log.add_paragraph("\n")
    save_verif_log(log)

    return df

# --- Tagging duplicates --- #
def tagging_duplicates(df, dups, variables):
    """
    This function tags duplicates in the dataframe based on specified variables
    :param df: The dataframe that the duplicates check should be performed on.
    :param variables: variables to check if there are duplicates.
    :return: df: The updated dataframe.
    """
    df[dups] = df.duplicated(subset=variables).astype(int)
    return df

# --- Printing summary statistics of all pwear variabels --- #
def pwear_statistics(log, df):
    """
    This function creates statistics for all Pwear variables and outputs them in the data qc log.
    :param log: The log book wear the verification checks gets printed to.
    :param df: The dataframe which is the imported summary means dataset.
    :return:
    """
    variables = [col for col in df.columns if col.startswith('Pwear')]

    # Adding to the qc_log
    paragraph = log.add_paragraph()
    run = paragraph.add_run("Compliance overview - Quadrant hours")
    run.bold=True

    # Adding table to qc log
    table = log.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.bold = True

    # Defining the header row
    hdr_cells = table.rows[0].cells
    headers = ['Variable', 'Count', 'Mean', '5th percentile', '25th percentile', '50th percentile', '75th percentile', '95th percentile']
    for i, header in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True

    # Formatting to get the statistics with only 2 decimals
    def format_statistics(stat):
        return f"{stat:.2f}" if pd.notnull(stat) else "N/A"

    # Loop through adding rows for each variable
    def add_table_row(table, data):
        row_cells = table.add_row().cells
        for i, value in enumerate(data):
            row_cells[i].text = value

    for variable in variables:
        data = df[variable]
        if data.empty:
            continue

    # Adding the statistics to the table
        stats = [
            variable,
            str(data.count()),
            format_statistics(data.mean()),
            format_statistics(data.quantile(0.05)),
            format_statistics(data.quantile(0.25)),
            format_statistics(data.quantile(0.50)),
            format_statistics(data.quantile(0.75)),
            format_statistics(data.quantile(0.95))
        ]
        add_table_row(table, stats)

    save_verif_log(log)


# --- Creating pwear_quad_diff variable to check that PWear is equal to the sum of Pwear of the quad times
def create_pwear_diff(df):
    """
    This function created variables in the dataframe that find the difference between pwear quadrants and wkend/wkday variables with the overall sum of these.
    :param df: The dataframe which is the imported summary means dataset.
    :return: df: The updated dataframe
    """
    df['Pwear_quad_diff'] = df['Pwear'] - (df['Pwear_morning'] + df['Pwear_noon'] + df['Pwear_afternoon'] + df['Pwear_night'])
    df['Pwear_wk_wkend_diff'] = df['Pwear'] - (df['Pwear_wkend'] + df['Pwear_wkday'])
    week_variables = ['wkend', 'wkday']

    for index, var in enumerate(week_variables):
        df[f'Pwear_{var}_quads_diff'] = df[f'Pwear_{var}'] - (df[f'Pwear_morning_{var}'] + df[f'Pwear_noon_{var}'] + df[f'Pwear_afternoon_{var}'] + df[f'Pwear_night_{var}'])
    return df

def proportion_categories(df):
    """ This function finds the differences in all variable different thresholds (so the difference between e.g., enmo_1plus and enmo_2plus and so on)
    :param df: The dataframe which is the imported summary means dataset.
    :return: The updated dataframe
    """
    if config.get('remove_thresholds').lower() == 'no':

        for var in config.get('verify_vars'):
            def loop_thresholds(a, b, c, df, var):
                for i in range(a, b, c):
                    l = i + c
                    df[f'{var}_prop_cat_{i}_{l}'] = df[f'{var}_{i}plus'] - df[f'{var}_{l}plus']

            loop_thresholds(0, 5, 1, summary_df, var)
            loop_thresholds(5, 150, 5, summary_df, var)
            loop_thresholds(150, 300, 10, summary_df, var)
            loop_thresholds(300, 1000, 100, summary_df, var)
            loop_thresholds(1000, 4000, 1000, summary_df, var)

            df[f'{var}_prop_cat_4000plus'] = df[f'{var}_4000plus']

            # Checking the total proportions of all categories equals the proportion of time spent above 0mg
            prop_cat_columns = [col for col in df.columns if col.startswith(f'{var}_prop_cat_')]
            df[f'{var}_total_prop'] = df[prop_cat_columns].sum(axis=1)
            df[f'{var}_total_diff'] = (df[f'{var}_total_prop'] - df[f'{var}_0plus']).abs()
    return df


# --- Printing verification checks to log --- #
def verif_checks(comparison_operator, variable, cut_off, df, log, text_to_log, column_number, list_of_headers, text_no_error):
    """
    This function compare the value of a defined variable to a cut_off value. If the variable is above/below a certain cut-off, a warning message will be printed to the data QC log and data for this file will be printed to the log.
    :param comparison_operator: Defining if the variable we are checking should be <, > or not equal to a cut-off value.
    :param variable: The variable from the dataframe that we are checking.
    :param cut_off: The cut-off that we are checking if a variable is <, > or not equal to.
    :param df: Dataframe with variables that are being checked.
    :param log: Vericiation log where the check are being outputted to.
    :param text_to_log: The text that we want to print to the verification log if some data is not how we expect (as of above conditions)
    :param column_number: The number of columns that we want printing to the table in the verification log.
    :param list_of_headers: Listing the variables that we want printed in the verification log in case the conditions above are met.
    :param x, y, z: Specifying the color that the text in the verification log is being printed in.
    :param text_no_error: The text that we want to print to the
    :return: None
    """
    # Defining mapping of string operators to their python function to call them in the function
    operators = {
        "<": operator.lt,
        ">": operator.gt,
        "!=": operator.ne,
        ">=": operator.ge,
        "<=": operator.le,
        "==": operator.eq
    }

    # Creating variables as a list that can store multiple variables
    if isinstance(variable, str):
        variable = [variable]

    # Checking if the length of the list 'variables' is 1 or 2 and then creating the condition that needs checking based on this:
    if len(variable) == 1:
        # Counting
        condition = operators[comparison_operator](df[variable[0]], cut_off)
    elif len(variable) == 2:
        condition1 = operators[comparison_operator](df[variable[0]], cut_off)
        condition2 = df[variable[1]].notna()
        condition = condition1 & condition2
    else:
        raise ValueError("This function only support 1 or 2 values")

    count = df[condition].shape[0]
    if count > 0:

        # Adding to the verif_log
        add_text(log, text_to_log, 255, 0, 0)

        # Adding table to verif log
        table = log.add_table(rows=1, cols=column_number)
        table.style = 'Table Grid'

        # Defining the header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(list_of_headers):
            run = hdr_cells[i].paragraphs[0].add_run(header)
            run.bold = True

        # Adding the data to the table:
        filtered_df = df.loc[condition]
        for _, row in filtered_df.iterrows():
            row_cells = table.add_row().cells
            for i, header in enumerate(list_of_headers):
                row_cells[i].text = str(row[header])

        log.add_paragraph("\n")
        save_verif_log(log)

    else:
        # Adding to the qc_log
        add_text_no_error(log, text_no_error)
        save_verif_log(log)


# --- LOOKING AT OUTLIERS IN DATA --- #
def outliers(df, log, list_variables, extra_variable, sort, text_to_log, filtering, level):
    """
    This function prints out sorted variables to the verification log to check for outliers.
    :param df: Dataframe with variables that are being checked.
    :param log: Vericiation log where the check are being outputted to.
    :param list_variables: list of variables that are being printed to the log to check for outliers.
    :param extra_variable: If REMOVE_THRESHOLD is NO this extra variable will be added to the list of variables that are being printed to the verification log.
    :param sort: Before printing the data to the log, the data will be sorted by this variable so it is easier to see the lowest and highest values.
    :param text_to_log: The text that we want to print to the verification log
    :return: None
    """
    # Adding variable to list of variables to print out
    if config.get('remove_thresholds').lower() == 'no' and extra_variable is not None:
        list_variables.append(extra_variable)


    # Adding text and table to log
    add_text(log, text_to_log, 0,0, 0)
    # calculating enmo mean for the dataset to print to verification log for reference
    if level == 'summary':
        enmo_mean = round(df['enmo_mean'].mean(), 2)
        add_text(log, f'Enmo mean for this dataset is {enmo_mean}. Use this value when going through the {filtering} values of enmo mean in the dataset.', 0, 0, 0)

    table = log.add_table(rows=1, cols=len(list_variables))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, var in enumerate(list_variables):
        run = hdr_cells[i].paragraphs[0].add_run(var)
        run.bold = True

    # Sorting dataframe by variable
    df = df.sort_values(by=sort)

    # Filtering to only print out 10 files with highest/lowest value of enmo_mean
    if len(df) > 10:
        if filtering == 'lowest':
            df = df.nsmallest(10, sort)
        if filtering == 'highest':
            df = df.nlargest(10, sort)

    # Adding data to table
    for index, row_data in df.iterrows():
        if row_data[sort] != -1 and not pd.isna(row_data[sort]):
            row_cells = table.add_row().cells
            for i, var in enumerate(list_variables):
                value = row_data[var]
                if isinstance(value, (int, float)):
                    value = round(value, 2)
                row_cells[i].text=str(value)
    log.add_paragraph("\n")
    save_verif_log(log)


# Add summary statistic headers to log:
def sum_stat_header(log):
    """
    Creating a table and printing headers to the table
    :param log: The verification log where the table is being outputted
    :return: table: Returning the table to be able to add the statistics afterwards.
    """
    table = log.add_table(rows=1, cols=9)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Variable', 'Count', 'Mean', 'Std', 'Min', '25%', '50%', '75%', 'Max']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    return table


# Add summary statistics to log:
def sum_stat_to_log(table, summary_stats, variable):
    """
    Printing the summary statistics to table. The sum_stat_header needs to be run prior to running this, to create the table
    :param table: The table returned from sum_stat_header
    :param summary_stats: The summary statistics created by summarising (describe()) one or multiple variables.
    :param variable: The variable(s) that is being summarised
    :return:
    """
    row_cells = table.add_row().cells

    # Adding overall statistics
    row_cells[0].text = variable
    row_cells[1].text = str(int(summary_stats['count']))
    row_cells[2].text = f"{summary_stats['mean']:.2f}"
    row_cells[3].text = f"{summary_stats['std']:.2f}"
    row_cells[4].text = f"{summary_stats['min']:.2f}"
    row_cells[5].text = f"{summary_stats['25%']:.2f}"
    row_cells[6].text = f"{summary_stats['50%']:.2f}"
    row_cells[7].text = f"{summary_stats['75%']:.2f}"
    row_cells[8].text = f"{summary_stats['max']:.2f}"


# --- Getting summary statistics for specified variable and printing to verification log --- #
def get_summary_stats(condition_operator, df, log, variables, text_to_log, description, text_no_files):
    """
    Summarising variables and printing summary statistics to log
    :param condition_operator: Condition of which the summary statistics count is being checked with
    :param df: Dataframe with the variables in
    :param log: The verification log where the statistics are being outputted
    :param variables: The variable(s) that is being summarised
    :param text_to_log: Text to print to log to describe what is being checked
    :param description: Discription to print to log to further explain what is being checked.
    :param text_no_files: Text to print to log if no files to summarise
    :return:
    """
    if config.get('remove_thresholds').lower() == 'no':

        operators = {
            "<": operator.lt,
            ">": operator.gt,
            "!=": operator.ne,
            ">=": operator.ge,
            "<=": operator.le
        }
        condition_func = operators.get(condition_operator)

        # Adding to the verification log
        add_text(log, text_to_log, 0, 0, 0)
        add_description_text(log, description, 0, 0, 0)

        # Adding table to verification log
        table = sum_stat_header(log)

        # Flag to indicate if any files found to be included
        flag_data_include = False

        # Getting the summary statistics for each variable
        for i, variable in enumerate(variables):
            summary_stats = df[variable].describe()
            if condition_func(summary_stats['count'], 0):
                flag_data_include = True

                # Adding overall statistics
                sum_stat_to_log(table, summary_stats, variable)

        # If no files to include, print message to log
        if not flag_data_include:
            add_text(log, text_no_files, 255, 0, 0)

        log.add_paragraph("\n")
        save_verif_log(log)


#--- Checking for any enmo negative values --- #
def check_negative_values(df, log, text_to_log, description, variables, text_no_error):
    '''
    This function checks if any files have negative values of any of the enmo variables. If any negative files it will print the first of these for each file.
    :param df: Dataframe that it is looking within for negative enmo values.
    :param log: The log where it is being printed to.
    :param text_to_log: Text to log explaining what it is showing.
    :param description: Text to further explain what to look for.
    :param variables: Variables that are being verified.
    :param text_no_error: Text to log if no files meets the criteria.
    :return:
    '''
    # Creating flag and list if any variables with negative values
    flag_negative_found = False
    files_with_negative_values = set()
    grouped = df.groupby('id')

    # Adding files with negative values to list
    for file_id, group in grouped:
        for _, row in group.iterrows():
            for var in variables:
                if row[var] < 0:
                    files_with_negative_values.add(file_id)
                    break

    # If any files with negative values, print it to log
    if files_with_negative_values:

        add_text(log, text_to_log, 0, 0, 0)
        add_description_text(log, description, 0, 0, 0)

        table = log.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        headers = ['File ID', 'Variable name', 'Variable value', 'Include value']
        for idx, header in enumerate(headers):
            run = hdr_cells[idx].paragraphs[0].add_run(header)
            run.bold = True

        include_present = 'include' in df.columns

        for file_id in files_with_negative_values:
            file_group = grouped.get_group(file_id)
            for _, row in file_group.iterrows():
                for var in variables:
                    if row[var] < 0:
                        value = row[var]
                        include_value = row['include'] if include_present else 'N/A'

                        row_cells = table.add_row().cells
                        row_cells[0].text = str(file_id)
                        row_cells[1].text = var
                        row_cells[2].text = str(value)
                        row_cells[3].text = str(include_value)
                        break
                break
        flag_negative_found = True

    # If no variables with negative values, printing a message to log.
    if not flag_negative_found:
        add_text_no_error(log, text_no_error)

    log.add_paragraph("\n")
    save_verif_log(log)

# --- Comparing ENMO_N and ENMO_0_99999 --- #
def compare_enmo(df, log, create_var, var_diff, text_to_log, text_no_error):
    """
    This function is comparing ENMO_n and ENMO_0_99999
    :param df: Dataframe with variables that are being checked.
    :param log: Verficiation log where the checks are being outputted
    :param create_var: Name of variable that is created by ENMO_0_plus * 720
    :param var_diff: The difference variable
    :param text_to_log: Text to print to log to indicate what check is being done
    :param text_no_error: Text in case there are no differences
    :return:
    """
    if config.get('remove_thresholds').lower() == 'no':
        if config['count_prefixes'].lower() == '1h':
            df[create_var] = df['ENMO_0plus'] * 720
        if config['count_prefixes'].lower() == '1m':
            df[create_var] = df['ENMO_0plus'] * 12

        # Generating difference between ENMO_n and ENMO_0plus_check
        df[var_diff] = df.apply(lambda x: abs(x['ENMO_n'] - x[create_var]) if x['ENMO_n'] != x[create_var] else 0, axis=1)

        # Summarising the difference
        sum_diff = df[df[var_diff] > 1][var_diff].describe()

        if sum_diff['count'] > 0:
            # Adding header to log and creating table with the summary statistics
            add_text(log, text_to_log, 0, 0, 0)
            table = sum_stat_header(log)
            sum_stat_to_log(table, sum_diff, var_diff)

        else:
            add_text_no_error(log, text_no_error)

        log.add_paragraph("\n")
        save_verif_log(log)

    return df


# --- COUNTING DIFFERENCES IN ENMO_N AND ENMO_OPLUS NOT DUE TO FIRST OR LAST TIMEPOINT IN THE FILE --- #
def diff_enmo(df, var1, var2, var3, log, text_to_log, text_no_error):
    """
    Formatting datetime variables and calculating differences between timepoint variables. Printing to log if there is any differences.
    :param df: Dataframe with the data.
    :param var1: timepoint variable --> first_file_timepoint
    :param var2: timepoint variable --> last_file_timepoint
    :param var3: timepoint variable --> DATETIME_ORIG
    :param log: Verification log.
    :param text_to_log: Text to print if any differences
    :param text_no_error: Text to print if no differences.
    :return:
    """
    # Converting timestamps to datetime format
    for var in (var1, var2, var3):
        df[var] = pd.to_datetime(df[var])

    # Generating differences between timepoint variables and DATETIME_ORIG
    for time_var in (var1, var2):
        df[f'diff_{time_var}'] = abs(df[time_var] - df[var3]) / pd.Timedelta(hours=1)

    if config.get('remove_thresholds').lower() == 'no':
        count = df[(df['ENMO_n_0plus_diff'] > 0.01) & (df[f'diff_{var1}'] > 1) & (df[f'diff_{var2}'] > 0)].shape[0]

        if count > 0:
            add_text(log, text_to_log, 255, 0, 0)
            add_text(log, f'Count: {count}', 0, 0, 0)

        else:
            add_text_no_error(log, text_no_error)

        log.add_paragraph("\n")
        save_verif_log(log)


# --- Summarising enmo variables --- #
def sum_enmo(remove_threshold, df, log, variables, text_to_log, description, text_no_files):
    """
    Summarising ENMO variables and printing to log
    :param df: Dataframe with enmo variables in
    :param log: verification log
    :param variables: variable that is being summarised
    :param text_to_log: Text to explain what is being checked
    :param text_no_files: Text if there are no files to summarize
    :return:
    """

    if config.get('remove_thresholds') == remove_threshold:

        # Adding to the verification log
        add_text(log, text_to_log, 0, 0, 0)
        add_description_text(log, description, 0, 0, 0)
        table = sum_stat_header(log)

        for i, variable in enumerate(variables):
            sum_stat = df[variable].describe()

            if sum_stat['count'] > 0:
                sum_stat_to_log(table, sum_stat, variable)

            else:
                add_text(log, text_no_files, 255, 0, 0)

        log.add_paragraph("\n")
        save_verif_log(log)

# --- CHECKING FILES/TIMEPOINT THAT ARE FLAGGED AS MECHANICAL NOISE --- #
def enmo_flag(df, log, variable, description, table_variables, text_no_error):
    """Counting how many files have an ENMO_mean above a specified flag
    :param df: Dataframe with the files.
    :param log: The verification log where the count it printed to.
    :param flag: The flag to count files with ENMO_mean above this
    :param variable: ENMO_mean
    :return: None
    """
    if config.get('processing').lower() == 'wave':
        count = df[(df[variable] == 1) & (df['QC_anomalies_total'] == 0)].shape[0]
    if config.get('processing').lower() == 'pampro':
        count = df[(df[variable] == 1)].shape[0]

    add_text(log, text_to_log=f"Number of timepoints in the data that are flagged as mechanical noise. \n Count: {count}", x=0, y=0, z=0)

    if count > 0:
        add_description_text(log, description, 0, 0, 0)
        filtered_df = df[df[variable]==1].sort_values(by='ENMO_mean')
        table = log.add_table(rows=1, cols=len(table_variables))
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        for i, var in enumerate(table_variables):
            run = hdr_cells[i].paragraphs[0].add_run(var)
            run.bold = True

        # Adding data to table
        for index, row_data in filtered_df.iterrows():
            row_cells = table.add_row().cells
            for i, var in enumerate(table_variables):
                value = row_data[var]
                if isinstance(value, (int, float)):
                    value = round(value, 2)
                row_cells[i].text = str(value)

    else:
        add_text_no_error(log, text_no_error)

    save_verif_log(log)

# Checking imputed variables
def verif_impute_data(df):
    df['include'] = 0

    # Checking if FLAG_NO_VALID_DAYS is a variable in the dataframe and otherwise it will give it the value 0
    FLAG_NO_VALID_DAYS_exists = 'FLAG_NO_VALID_DAYS' in df.columns
    FLAG_NO_VALID_DAYS_condition = (df['FLAG_NO_VALID_DAYS'] != 1) if FLAG_NO_VALID_DAYS_exists else True
    CALIBRATION_TYPE_exists = 'calibration_type' in df.columns
    CALIBRATION_TYPE_condition = (df['calibration_type'] != 'fail') if CALIBRATION_TYPE_exists else True

    # include = 1 = Pwear and Pwear for all quadrants above threshold.
    df.loc[
        (df['Pwear'] >= config.get('ver_pwear')) &
        (df['Pwear'].notnull()) &
        (df['Pwear_morning'] >= config.get('ver_pwear_morn')) &
        (df['Pwear_noon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_afternoon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_night'] >= config.get('ver_pwear_quad')) &
        (CALIBRATION_TYPE_condition) &
        (FLAG_NO_VALID_DAYS_condition), 'include'] = 1

    # 2 = Pwear above threshold and pwear for noon, afternoon and night above threshold (morning under, so we want to impute sleep)
    df.loc[
        (df['Pwear'] >= config.get('ver_pwear')) &
        (df['Pwear'].notnull()) &
        (df['Pwear_morning'] < config.get('ver_pwear_morn')) &
        (df['Pwear_noon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_afternoon'] >= config.get('ver_pwear_quad')) &
        (df['Pwear_night'] >= config.get('ver_pwear_quad')) &
        (CALIBRATION_TYPE_condition) &
        (FLAG_NO_VALID_DAYS_condition) &
        (df['include'] != 1), 'include'] = 2

    df['imputed'] = 0
    # imputed = 2 = Data imputed between 00:00-06:00 if not worn (proportional to Pwear)
    df.loc[(df['include'] == 2), 'imputed' ] = 1

    # Generating sedentary, light and mvpa variables
    types = ['', '_IMP']
    for type_ in types:
        for x in range(25, 40, 5):
            column_name = f'sed_{x}{type_}'
            df[column_name] = (1 - df[f'enmo_{x}plus{type_}']) * 1400

        light_name = f'lpa{type_}'
        df[light_name] = (df[f'enmo_30plus{type_}'] - df[f'enmo_125plus{type_}']) * 1440

        for x in range(100, 175, 25):
            mvpa_name = f'mvpa_{x}{type_}'
            df[mvpa_name] = df[f'enmo_{x}plus{type_}'] * 1400

    return df


def print_impute_checks(df, log, text_to_log, description, variables, condition_operator, text_no_files):
    operators = {
        "<": operator.lt,
        ">": operator.gt,
        "!=": operator.ne,
        ">=": operator.ge,
        "<=": operator.le
    }
    condition_func = operators.get(condition_operator)

    # Adding to the verification log
    add_text(log, text_to_log, 0, 0, 0)
    add_description_text(log, description, 0, 0, 0)

    # Adding table to verification log
    table = sum_stat_header(log)

    # Flag to indicate if any files found to be included
    flag_data_include = False

    # Getting the summary statistics for each variable
    for i, variable in enumerate(variables):
        summary_stats = df[variable].describe()
        if condition_func(summary_stats['count'], 0):
            flag_data_include = True

            # Adding overall statistics
            sum_stat_to_log(table, summary_stats, variable)

    # If no files to include, print message to log
    if not flag_data_include:
        add_text(log, text_no_files, 255, 0, 0)

    log.add_paragraph("\n")
    save_verif_log(log)



if __name__ == '__main__':

    if len(sys.argv) < 2:
        print("Error: No config file provided.")
        sys.exit(1)

    config_path = sys.argv[1]
    config = load_config(config_path)

    print("Loaded config:", config)
    # Now you can use config values inside your script
    config["pc_date"] = datetime.utcnow().isoformat()

    # --- SECTION 1: VERIFICATION OF OUTPUT SUMMARY OVERALL MEANS --- #
    # Creating verification log and importing summary dataframe
    verif_log = create_verif_log("VERIFICATION LOG")
    summary_df, summary_file_exists = dataframe(file_name=config.get('sum_output_file'), variable='id')

    # If dataframe exists, print out files processed, devices used and summary of start dates
    if summary_file_exists:
        add_header(log_header=f"VERIFICATION OF OUTPUT SUMMARY OVERALL MEANS")        
        information_to_verif_log(log=verif_log, df=summary_df, table='No', variable_to_count='id', text_to_log="Number of files processed", count_mode='total')
        information_to_verif_log(log=verif_log, df=summary_df, table='Yes', variable_to_count='device', text_to_log="Number of devices used", count_mode='unique')
        df = sum_startdate(log=verif_log, df=summary_df, text_to_log="Summary of start dates:", description="Check that the minimum and maximum start date falls within the expected testing dates.", x=0, y=0, z=0)

        # Check files with no valid wear time or wear log info (if relevant)
        if 'FLAG_NO_VALID_DAYS' in df.columns:
            column_number_arg = 6 if config.get('use_wear_log').lower() == 'yes' else 4
            list_of_headers_arg = ['id', 'startdate', 'Pwear', 'enmo_mean', 'start', 'end'] if config.get('use_wear_log').lower() == 'yes' else ['id', 'startdate', 'Pwear', 'enmo_mean']
            verif_checks(comparison_operator="==", variable="FLAG_NO_VALID_DAYS", cut_off=1, df=summary_df, log=verif_log, text_to_log="Some files did not have any valid days.", column_number=column_number_arg,
                         list_of_headers=list_of_headers_arg, text_no_error="All files have valid wear time. No ID's to check.")

        # Create include variable
        df = include_criteria(log=verif_log, df=summary_df, text_to_log="Summarising include variable:", description= f"Include = 0: Include criteria are not met (Pwear < {config.get('ver_pwear')}, Pwear morning < {config.get('ver_pwear_morn')} and Pwear noon/afternoon/night < {config.get('ver_pwear_quad')}). \n Include = 1: Include criteria are met (Pwear >= {config.get('ver_pwear')}, Pwear morning >= {config.get('ver_pwear_morn')} and Pwear noon/afternoon/night >= {config.get('ver_pwear_quad')}).")

        # Tagging duplicates
        df = tagging_duplicates(df=summary_df, dups='duplicates_data', variables=['device', 'generic_first_timestamp', 'generic_last_timestamp'])
        df = tagging_duplicates(df=summary_df, dups='duplicates_id', variables=['id'])

        # Dropping duplicates that have already been investigated
        if config.get('run_housekeeping').lower() == 'yes':
            summary_df = summary_df[(~summary_df['id'].isin(filenames_to_remove))]

        # Printing files that have not calibrated
        verif_checks(
            comparison_operator=">",
            variable="file_end_error",
            cut_off=config.get('cal_error'),
            df=summary_df,
            log=verif_log,
            text_to_log=f"Some files have not calibrated and have an end error above {config.get('cal_error')}- check file end errors. \n Uncalibrated data will be set to missing during post-processing.",
            column_number=3,
            list_of_headers=['id', 'file_start_error', 'file_end_error'],
            text_no_error=f"All files have calibrated. No files had an end error above {config.get('cal_error')}, no files to check.")

        if config.get('processing').lower() == 'pampro':
            verif_checks(
                comparison_operator=">",
                variable="mf_end_error",
                cut_off=config.get('cal_error'),
                df=summary_df,
                log=verif_log,
                text_to_log=f"Some files have not calibrated and have an end error above {config.get('cal_error')}  - check mf file end errors. \n Uncalibrated data will be set to missing during post-processing.",
                column_number=3,
                list_of_headers=['id', 'mf_start_error', 'mf_end_error'],
                text_no_error=f"All files have calibrated. No files had an mf end error above {config.get('cal_error')}, no files to check.")

        # Printing out duplicates
        verif_checks(
            comparison_operator="!=",
            variable="duplicates_data",
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are duplicates in this summary dataset. \n Add the duplicate file to the Housekeeping file to remove data from final dataset.",
            column_number=4,
            list_of_headers=['id', 'device', 'generic_first_timestamp', 'generic_last_timestamp'],
            text_no_error="There are no duplicated data in this summary dataset")

        verif_checks(
            comparison_operator="!=",
            variable="duplicates_id",
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are files with duplicate IDs. \n Add the duplicate file to the Housekeeping file to remove data from final dataset.",
            column_number=3,
            list_of_headers=['id', 'device', 'generic_first_timestamp'],
            text_no_error="There are no duplicates based on id only."
        )

        # Checking length of recording
        verif_checks(
            comparison_operator="<",
            variable='RecordLength',
            cut_off=config.get('min_inclusion_hrs'),
            df=summary_df,
            log=verif_log,
            text_to_log=f"There are potential battery functionality issues. Some recordings are less than {config.get('min_inclusion_hrs')} hours. It is recommended to test the following devices.",
            column_number=3,
            list_of_headers=['id', 'device', 'RecordLength'],
            text_no_error=f"All recordings are above {config.get('min_inclusion_hrs')} hours. No devices to check."
        )

        # Checking for any anomalies
        variable_arg = ['QC_anomaly_G', 'QC_anomaly_F'] if config.get('processing').lower() == 'wave' else ['Anom_F']
        anomaly_column_number_arg = 4 if config.get('processing').lower() == 'wave' else 3
        anomaly_list_headers_arg = ['id', 'device', 'QC_anomaly_G', 'QC_anomaly_F'] if config.get('processing').lower() == 'wave' else ['id', 'device', 'Anom_F']
        verif_checks(
            comparison_operator=">",
            variable=variable_arg,
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are timestamp anomaly F's or G's in the files. It is recommended to remove these devices from circulation.",
            column_number=anomaly_column_number_arg,
            list_of_headers=anomaly_list_headers_arg,
            text_no_error="There are no timestamp anomalies in the files. No files to check.")

        # Checking first and last battery percentage
        verif_checks(
            comparison_operator="<",
            variable='qc_first_battery_pct',
            cut_off=75,
            df=summary_df,
            log=verif_log,
            text_to_log="Some devices were set up with low battery (<75%)",
            column_number=4,
            list_of_headers=['id', 'device', 'qc_first_battery_pct', 'qc_last_battery_pct'],
            text_no_error="All devices had >75% battery when set up.")

        verif_checks(
            comparison_operator="<",
            variable='qc_last_battery_pct',
            cut_off=10,
            df=summary_df,
            log=verif_log,
            text_to_log="Some devices had a low battery percentage at the end of recording (<10%).",
            column_number=4,
            list_of_headers=['id', 'device', 'qc_first_battery_pct', 'qc_last_battery_pct'],
            text_no_error="All devices had >10% battery at the end of recording. No devices to check.")

        # Checking frequency devices was set up with
        verif_checks(
            comparison_operator="!=",
            variable='frequency',
            cut_off=config.get('protocol_frequency'),
            df=summary_df,
            log=verif_log,
            text_to_log=f"There are files not initialised at the protocol frequency of {config.get('protocol_frequency')} Hz",
            column_number=3,
            list_of_headers=['id', 'device', 'frequency'],
            text_no_error=f"All devices were initialised at the protocol frequency of {config.get('protocol_frequency')} Hz."
        )


        # Printing summary statistics of all pwear variables
        landscape(verif_log)
        pwear_statistics(log=verif_log, df=summary_df)
        portrait(verif_log)

        # Printing out files with negative values of ENMO_mean
        verif_checks(
            comparison_operator="<",
            variable='enmo_mean',
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are negative values of ENMO_mean in this summary dataset.",
            column_number=5,
            list_of_headers=['id', 'enmo_mean', 'Pwear', 'RecordLength', 'include'],
            text_no_error="There are no negative values of ENMO_mean in this summary dataset. No IDs to check.")

        # Printing out data to look for outliers
        outliers(summary_df, verif_log, list_variables = ['id', 'enmo_mean', 'Pwear', 'RecordLength'], extra_variable='enmo_0plus', sort='enmo_mean',
                 text_to_log="Look through the lowest values of enmo_mean for potential outliers.", filtering='lowest', level='summary')
        outliers(summary_df, verif_log, list_variables = ['id', 'enmo_mean', 'Pwear', 'RecordLength'], extra_variable='enmo_0plus', sort='enmo_mean',
                 text_to_log="Look through the highest values of enmo_mean for potential outliers.", filtering='highest', level='summary')

        # Creating pwear difference variables and checking if the quadrants, weekend/wkday is equal to overall pwear
        df = create_pwear_diff(summary_df)
        verif_checks(
            comparison_operator=">=",
            variable=['Pwear_quad_diff', 'Pwear'],
            cut_off=0.0001,
            df=summary_df,
            log=verif_log,
            text_to_log="The sum of Pwear is not equal to the sum of Pwear for all quadrants. Look through the files below:",
            column_number=3,
            list_of_headers=['id', 'Pwear', 'Pwear_quad_diff'],
            text_no_error="The Pwear overall sum is equal to the sum of Pwear for all quadrants. No IDs to check.")
        verif_checks(
            comparison_operator=">=",
            variable=['Pwear_wk_wkend_diff', 'Pwear'],
            cut_off=0.0001,
            df=summary_df,
            log=verif_log,
            text_to_log="The sum of Pwear is not equal to the sum of Pwear for weekend and weekdays. Look through the files below:",
            column_number=3,
            list_of_headers=['id', 'Pwear', 'Pwear_wk_wkend_diff'],
            text_no_error="The sum of Pwear is equal to the sum of Pwear for weekday and weekends. No IDs to check.")
        weekdays = ['wkday', 'wkend']
        for var in weekdays:
            verif_checks(
                comparison_operator=">=",
                variable=[f'Pwear_{var}_quads_diff', 'Pwear'],
                cut_off=0.0001,
                df=summary_df,
                log=verif_log,
                text_to_log=f"The sum of Pwear_{var} is not equal to the sum of Pwear_{var} for all quadrants. Look through the files below:",
                column_number=3,
                list_of_headers=['id', f'Pwear_{var}', f'Pwear_{var}_quads_diff'],
                text_no_error=f" The sum of Pwear_{var} is equal to the sum of Pwear_{var} for all quadrants. No IDs to check.")

        # Creating difference variables between all thresholds and checking that the total proportion of all categories equals the proportion of time spent above 0mg
        df = proportion_categories(summary_df)
        for verif_var in config.get('verify_vars'):
            verif_checks(
                comparison_operator=">=",
                variable=[f'{verif_var}_total_diff', 'Pwear'],
                cut_off=0.0001,
                df=summary_df,
                log=verif_log,
                text_to_log=f"The sum of {verif_var} proportions does not equal the total proportion. Look through the files below:",
                column_number=3,
                list_of_headers=['id', f'{verif_var}_0plus', f'{verif_var}_total_prop'],
                text_no_error=f"The sum of {verif_var} proportions is equal to the total proportion for all files. No IDs to check.")

        # Getting summary statistics for ENMO variables (overall and only on data that meets the inclusion criteria)
        get_summary_stats(condition_operator="!=", df=summary_df, log=verif_log, variables=['enmo_0plus'], text_to_log="Overall summary statistics for enmo_0plus", description="enmo_0plus is the proportion of time spent above >= 0 milli-g. This should be ~1. This indicates how much the device has been worn for.",
 text_no_files="No observations to summarize")
        get_summary_stats(condition_operator="!=", df=summary_df[summary_df['include'] == 1], log=verif_log, variables=['enmo_0plus'], text_to_log="Summary statistics for enmo_0plus where include=1", description="enmo_0plus is the proportion of time spent above >= 0 milli-g. This should be ~1. This indicates how much the device has been worn for.", text_no_files="No observations (include == 1) to summarize")
        enmo_variables = [col for col in summary_df.columns if col.startswith('enmo_') and col.endswith('plus')]
        check_negative_values(df=summary_df, log=verif_log, text_to_log="There are negative values in the enmo_*plus variables.", description="Check to see if device has calibrated correctly. \n It is suggested to remove data/file if any negative values are present", variables= enmo_variables, text_no_error="There are no files with negative values in any of the enmo_variables. No files to check.")
        check_negative_values(df=summary_df[summary_df['include'] == 1], log=verif_log, text_to_log="There are negative values in the enmo_*plus variables (and where include=1).", description="Check to see if device has calibrated correctly. \n It is suggested to remove data/file if any negative values are present", variables= enmo_variables, text_no_error="There are no files with negative values in any of the enmo_variables (and where include == 1). No files to check.")

        # Printing out of some files were unable to process (no_analysis_files)
        if 'flag_unable_to_process' in summary_df.columns:
            verif_checks(comparison_operator='==', variable='flag_unable_to_process', cut_off=1, df=summary_df, log=verif_log, text_to_log='Some files were unable to process (they did not have an hourly/minute level file). \nOnly metadata for this file is included in the release.',
                     column_number=3, list_of_headers=['id', 'file_start_error', 'file_end_error'], text_no_error='All files were able to process. No files to check.')

        # IMPUTED VARIABLE CHECKING - DOES IT MATCH FULLY WORN FILES WITH ADDITIONAL SLEEP ADDED

        if config.get('impute_data').lower() == 'yes':
            summary_df = verif_impute_data(df=summary_df)

            print_impute_checks(df=summary_df[summary_df['include'] == 1], log=verif_log, text_to_log="Check imputed variables - check how much imputing variables effects enmo_mean",
                              description="include = 1: Pwear overall and Pwear for each quadrant are above criteria. Enmo and IMP variables should not differ:", variables=['enmo_mean', 'enmo_mean_IMP', 'sed_30', 'sed_30_IMP', 'lpa', 'lpa_IMP', 'mvpa_125', 'mvpa_125_IMP'],
                              condition_operator="!=", text_no_files="There are no imputed variables to check.")

            print_impute_checks(df=summary_df[(summary_df['include'] == 2) & (summary_df['Pwear'] > 18)], log=verif_log, text_to_log="Check imputed variables - check how much imputing variables effects enmo_mean",
                              description="include = 2: Pwear overall and Pwear for 3 quadrants are above criteria.", variables=['enmo_mean', 'enmo_mean_IMP', 'sed_30', 'sed_30_IMP', 'lpa', 'lpa_IMP', 'mvpa_125', 'mvpa_125_IMP'],
                              condition_operator="!=", text_no_files="There are no imputed variables to check.")


    # --- SECTION 2: VERIFICATION OF HOURLY FILE(S) --- #
    # Importing hourly dataframe
    hourly_df, hourly_file_exists = dataframe(file_name=config.get('hour_output_file'), variable='file_id')
    add_header(log_header="VERIFICATION OF HOURLY FILES")

    # If dataframe exists, tag duplicates and print to log if there are any
    if hourly_file_exists:

        # Dropping duplicates that have already been investigated
        if config.get('run_housekeeping').lower() == 'yes':
            hourly_df = hourly_df[(~hourly_df['file_id'].isin(filenames_to_remove))]

        # Tagging duplicates
        tagging_duplicates_arg = ['timestamp', 'ENMO_mean', 'ENMO_30plus', 'ENMO_125plus']
        if 'pitch_mean' in hourly_df.columns:
            tagging_duplicates_arg.append('PITCH_mean')
        if 'roll_mean' in hourly_df.columns:
            tagging_duplicates_arg.append('ROLL_mean')
        df = tagging_duplicates(df=hourly_df, dups='dup_enmo_date', variables=tagging_duplicates_arg)

        # Printing out duplicates
        verif_checks(
            comparison_operator="!=",
            variable="dup_enmo_date",
            cut_off=0,
            df=hourly_df,
            log=verif_log,
            text_to_log="There are duplicates in this hourly dataset. \n Add the duplicate file to the Housekeeping file to remove data from final dataset.",
            column_number=5,
            list_of_headers=['id', 'file_id', 'device', 'timestamp', 'ENMO_mean'],
            text_no_error="There are no duplicated data in this hourly dataset")


        # Comparing ENMO_n and ENMO_0plus * 720. If these are not equal they will be printed to the log.
        df = compare_enmo(
            df=hourly_df,
            log=verif_log,
            create_var="ENMO_0plus_check",
            var_diff="ENMO_n_0plus_diff",
            text_to_log="Summarising the difference between ENMO_n and ENMO_0plus * 720, if these are not equal to each other",
            text_no_error="ENMO_n and ENMO_0plus * 720 are equal to each other for all observations. No IDs to check.")

        # Summarising ENMO_0plus if thresholds are not removed.
        sum_enmo(
            remove_threshold= "No",
            df=hourly_df,
            log=verif_log,
            variables=['ENMO_0plus'],
            text_to_log="Overall summary statistics for enmo_0plus.",
            description="enmo_0plus is the proportion of time spent above >= 0 milli-g. This should be ~1. This indicates how much the device has been worn for.",
            text_no_files="No observations to summarize.")


        # Summarising all ENMO variables if thresholds are not removed
        ENMO_variables = [col for col in hourly_df.columns if col.startswith('ENMO_') and col.endswith('plus')]
        check_negative_values(df=hourly_df, log=verif_log, text_to_log="There are negative values in the enmo_*plus variables.", description="Check to see if device has calibrated correctly. \n It is suggested to remove data/file if any negative values are present", variables= ENMO_variables, text_no_error="There are no files with negative values in any of the enmo_variables. No files to check.")

        # Printing out data sorted by ENMO mean to look through for potential outliers
        list_variables_arg = ['file_id', 'DATETIME_ORIG', 'ENMO_mean', 'ENMO_n', 'ENMO_missing', 'ENMO_sum', 'QC_anomalies_total', 'FLAG_MECH_NOISE'] if config.get('processing').lower() == 'wave' else ['file_id', 'DATETIME_ORIG', 'ENMO_mean', 'ENMO_n', 'ENMO_missing', 'ENMO_sum', 'FLAG_MECH_NOISE']
        landscape(verif_log)
        outliers(df=hourly_df, log=verif_log, list_variables=list_variables_arg, extra_variable=None, sort='ENMO_mean',
                 text_to_log="Look through the lowest values of enmo_mean for potential outliers.", filtering='lowest', level='hourly')
        outliers(df=hourly_df, log=verif_log, list_variables=list_variables_arg, extra_variable=None, sort='ENMO_mean',
                 text_to_log="Look through the highest values of enmo_mean for potential outliers.", filtering='highest', level='hourly')
        portrait(verif_log)

       # Checking files/timepoints that are flagged as mechanical noise
        if config['count_prefixes'].lower() == '1h':
            variables_table = ['id', 'dayofweek', 'hourofday', 'ENMO_mean', 'Pwear']
        if config['count_prefixes'].lower() == '1m':
            variables_table = ['id', 'dayofweek', 'hourofday', 'minuteofhour', 'ENMO_mean', 'Pwear']
            
        enmo_flag(
            df=hourly_df,
            log=verif_log,
            variable='FLAG_MECH_NOISE',
            description=
        """The files below have a high enmo_mean, which could suggest that there is mechanical noise in the data. \nDatapoints are flagged as mechanical noise if one of the following was applicable:
            - enmo_mean > 3000 and Pwear not 0. This would suggest an epoch either from monitor issue/other sources (i.e. washing machines)
            - enmo_mean > 1500 and pwear > 0.9. This would suggest that device was worn but that there is noise.
            - enmo_mean > 600 and pwear < 0.1. This would be an epochs that have a very small amount of Pwear but is getting a large amount of enmo for that blip of data \nIt is suggested to remove these hours from the hourly release file. The hours will be flagged in the release file (FLAG_MECH_NOISE=1).""",
            table_variables = ['id', 'dayofweek', 'hourofday', 'ENMO_mean', 'Pwear'],
            text_no_error='There are no timepoints flagged as mechanical noise. No files to check.')



